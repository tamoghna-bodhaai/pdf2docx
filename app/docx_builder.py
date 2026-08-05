"""Render the transcribed Markdown into an editable .docx document."""

from __future__ import annotations

import re
from io import BytesIO
from pathlib import Path

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import parse_xml
from docx.shared import Emu, Pt, RGBColor

from .latex_omml import display_math_xml, inline_math_xml
from .xml_text import xml_safe

# --------------------------------------------------------------------------- #
# Block-level patterns
# --------------------------------------------------------------------------- #

HEADING_RE = re.compile(r"^(?P<hashes>#{1,6})\s+(?P<text>.*)$")
FENCE_RE = re.compile(r"^\s*(?P<ticks>`{3,}|~{3,})\s*(?P<lang>[\w+-]*)\s*$")
HRULE_RE = re.compile(r"^\s*(?:\*\s*\*\s*\*[\s*]*|-\s*-\s*-[\s-]*|_\s*_\s*_[\s_]*)$")
BULLET_RE = re.compile(r"^(?P<indent>\s*)(?P<marker>[-*+])\s+(?P<text>.*)$")
ORDERED_RE = re.compile(r"^(?P<indent>\s*)(?P<marker>\d+)[.)]\s+(?P<text>.*)$")
QUOTE_RE = re.compile(r"^\s*>\s?(?P<text>.*)$")
TABLE_ROW_RE = re.compile(r"^\s*\|.*\|\s*$")
TABLE_SEP_RE = re.compile(r"^\s*\|(?:\s*:?-{2,}:?\s*\|)+\s*$")
DISPLAY_MATH_RE = re.compile(r"^\s*\$\$(?P<body>.*?)\$\$\s*$", re.S)
IMAGE_RE = re.compile(r"^\s*!\[(?P<alt>[^\]]*)\]\((?P<path>[^)]+)\)\s*$")

# Inline patterns, matched left-to-right by a single alternation.
INLINE_RE = re.compile(
    r"(?P<codefence>`+)(?P<code>.+?)(?P=codefence)"
    r"|(?<!\\)\$(?P<math>(?:\\.|[^$\\])+?)(?<!\\)\$"
    r"|\*\*(?P<bold>.+?)\*\*"
    r"|__(?P<bold2>.+?)__"
    r"|(?<!\w)\*(?P<italic>[^*\n]+?)\*(?!\w)"
    r"|(?<!\w)_(?P<italic2>[^_\n]+?)_(?!\w)"
    r"|~~(?P<strike>.+?)~~"
    r"|\[(?P<linktext>[^\]\n]*)\]\((?P<linkurl>[^)\n]*)\)",
    re.S,
)


def normalise_math_delimiters(text: str) -> str:
    r"""Rewrite \(..\) and \[..\] into $..$ and $$..$$ so one path handles math."""
    text = re.sub(r"\\\[(.+?)\\\]", lambda m: f"\n$$\n{m.group(1).strip()}\n$$\n", text, flags=re.S)
    text = re.sub(r"\\\((.+?)\\\)", lambda m: f"${m.group(1).strip()}$", text, flags=re.S)
    return text


# --------------------------------------------------------------------------- #
# Document scaffolding
# --------------------------------------------------------------------------- #


def _shading(fill: str) -> str:
    return (
        '<w:shd xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" '
        f'w:val="clear" w:color="auto" w:fill="{fill}"/>'
    )


def ensure_styles(document: Document) -> None:
    styles = document.styles

    if "Code Block" not in styles:
        style = styles.add_style("Code Block", WD_STYLE_TYPE.PARAGRAPH)
        style.base_style = styles["No Spacing"]
        style.font.name = "Consolas"
        style.font.size = Pt(9.5)
        paragraph_format = style.paragraph_format
        paragraph_format.left_indent = Pt(12)
        paragraph_format.space_before = Pt(6)
        paragraph_format.space_after = Pt(6)
        style.element.get_or_add_pPr().append(parse_xml(_shading("F4F5F7")))

    if "Code Inline" not in styles:
        style = styles.add_style("Code Inline", WD_STYLE_TYPE.CHARACTER)
        style.font.name = "Consolas"
        style.font.size = Pt(9.5)
        style.font.color.rgb = RGBColor(0xB0, 0x30, 0x60)

    if "Figure Note" not in styles:
        style = styles.add_style("Figure Note", WD_STYLE_TYPE.PARAGRAPH)
        style.base_style = styles["Normal"]
        style.font.italic = True
        style.font.size = Pt(9.5)
        style.font.color.rgb = RGBColor(0x55, 0x5D, 0x6B)
        style.paragraph_format.left_indent = Pt(18)


# --------------------------------------------------------------------------- #
# Inline rendering
# --------------------------------------------------------------------------- #


def _add_text_run(paragraph, text: str, *, bold=False, italic=False, strike=False):
    if not text:
        return
    run = paragraph.add_run(text)
    run.bold = bold or None
    run.italic = italic or None
    run.font.strike = strike or None


def render_inline(paragraph, text: str, *, bold=False, italic=False, strike=False) -> None:
    """Write `text` into `paragraph`, expanding inline Markdown and $math$."""
    text = text.replace("\\$", "\x00")
    position = 0
    for match in INLINE_RE.finditer(text):
        _add_text_run(
            paragraph,
            text[position : match.start()].replace("\x00", "$"),
            bold=bold, italic=italic, strike=strike,
        )
        position = match.end()

        if match.group("code") is not None:
            run = paragraph.add_run(match.group("code").replace("\x00", "$"))
            run.style = paragraph.part.document.styles["Code Inline"]
        elif match.group("math") is not None:
            latex = match.group("math").replace("\x00", "$")
            paragraph._p.append(parse_xml(inline_math_xml(latex)))
        elif match.group("bold") is not None or match.group("bold2") is not None:
            inner = match.group("bold") or match.group("bold2")
            render_inline(paragraph, inner, bold=True, italic=italic, strike=strike)
        elif match.group("italic") is not None or match.group("italic2") is not None:
            inner = match.group("italic") or match.group("italic2")
            render_inline(paragraph, inner, bold=bold, italic=True, strike=strike)
        elif match.group("strike") is not None:
            render_inline(paragraph, match.group("strike"), bold=bold, italic=italic, strike=True)
        elif match.group("linktext") is not None:
            label = match.group("linktext") or match.group("linkurl")
            render_inline(paragraph, label, bold=bold, italic=True, strike=strike)

    _add_text_run(
        paragraph,
        text[position:].replace("\x00", "$"),
        bold=bold, italic=italic, strike=strike,
    )


# --------------------------------------------------------------------------- #
# Block rendering
# --------------------------------------------------------------------------- #


def _list_style(document: Document, ordered: bool, level: int) -> str:
    base = "List Number" if ordered else "List Bullet"
    level = max(0, min(level, 2))
    name = base if level == 0 else f"{base} {level + 1}"
    return name if name in document.styles else base


def rule_border(paragraph) -> None:
    """Draw a horizontal rule as the bottom border of an empty paragraph."""
    paragraph.paragraph_format.space_before = Pt(2)
    paragraph._p.get_or_add_pPr().append(
        parse_xml(
            '<w:pBdr xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
            '<w:bottom w:val="single" w:sz="6" w:space="1" w:color="C7CBD1"/></w:pBdr>'
        )
    )


def place_picture(paragraph, source, width_pt: float, height_pt: float | None = None) -> bool:
    """Add a picture at its natural size, shrunk to fit the space available.

    Scaled down but never up: a figure cut out of a page is already the size the
    page drew it at, and stretching a 200-pixel diagram across the column would
    only make its own scan noise larger.
    """
    if isinstance(source, (bytes, bytearray)):
        source = BytesIO(bytes(source))
    try:
        picture = paragraph.add_run().add_picture(source)
    except Exception:
        return False  # an unreadable crop must not take the document down
    scale = 1.0
    if picture.width > Pt(width_pt):
        scale = Pt(width_pt) / picture.width
    if height_pt and picture.height * scale > Pt(height_pt):
        scale = Pt(height_pt) / picture.height
    if scale < 1.0:
        picture.width = int(picture.width * scale)
        picture.height = int(picture.height * scale)
    return True


class DocxWriter:
    def __init__(self, title: str | None = None) -> None:
        self.document = Document()
        ensure_styles(self.document)
        normal = self.document.styles["Normal"]
        normal.font.name = "Calibri"
        normal.font.size = Pt(11)
        normal.paragraph_format.space_after = Pt(8)
        if title:
            self.document.add_heading(xml_safe(title), level=0)

    # -- primitives -------------------------------------------------------- #

    def add_paragraph_text(self, text: str, style: str | None = None):
        paragraph = self.document.add_paragraph(style=style)
        render_inline(paragraph, text)
        return paragraph

    def add_display_math(self, latex: str) -> None:
        paragraph = self.document.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph._p.append(parse_xml(display_math_xml(latex)))

    def add_heading(self, text: str, level: int) -> None:
        style = f"Heading {min(max(level, 1), 6)}"
        paragraph = self.document.add_paragraph(
            style=style if style in self.document.styles else None
        )
        render_inline(paragraph, text)

    def add_rule(self) -> None:
        rule_border(self.document.add_paragraph())

    def add_list_item(self, text: str, marker: str, ordered: bool, level: int) -> None:
        paragraph = self.document.add_paragraph(style=_list_style(self.document, ordered, level))
        render_inline(paragraph, text)

    def add_quote(self, text: str) -> None:
        style = "Figure Note" if text.lstrip().startswith(("_Figure", "*Figure")) else "Quote"
        if style not in self.document.styles:
            style = "Figure Note"
        self.add_paragraph_text(text, style=style)

    @property
    def column(self) -> float:
        """Width of the text column, in points."""
        section = self.document.sections[0]
        usable = section.page_width - section.left_margin - section.right_margin
        return max(Emu(usable).pt, 72.0)

    def add_picture(self, path: str, caption: str = "") -> None:
        section = self.document.sections[0]
        # A figure taller than the page it is on pushes everything after it onto
        # the next one and leaves the page it came from empty.
        usable = section.page_height - section.top_margin - section.bottom_margin
        height = max(Emu(usable).pt, 72.0)
        paragraph = self.document.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if not place_picture(paragraph, path, self.column, height * 0.85):
            return
        if caption:
            self.add_quote(f"_{caption}_")

    def add_code_block(self, lines: list[str]) -> None:
        for line in lines or [""]:
            paragraph = self.document.add_paragraph(style="Code Block")
            paragraph.add_run(line)

    def add_table(self, rows: list[list[str]], alignments: list[str]) -> None:
        columns = max(len(row) for row in rows)
        table = self.document.add_table(rows=0, cols=columns)
        table.style = "Table Grid" if "Table Grid" in self.document.styles else "Normal Table"
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        for index, cells in enumerate(rows):
            row = table.add_row()
            for column in range(columns):
                cell = row.cells[column]
                cell.text = ""
                paragraph = cell.paragraphs[0]
                if column < len(alignments):
                    paragraph.alignment = {
                        "center": WD_ALIGN_PARAGRAPH.CENTER,
                        "right": WD_ALIGN_PARAGRAPH.RIGHT,
                    }.get(alignments[column], WD_ALIGN_PARAGRAPH.LEFT)
                value = cells[column] if column < len(cells) else ""
                render_inline(paragraph, value, bold=(index == 0))

    def add_page_break(self) -> None:
        self.document.add_page_break()

    def save(self, path: Path) -> None:
        self.document.save(path)


# --------------------------------------------------------------------------- #
# Markdown driver
# --------------------------------------------------------------------------- #


def _split_table_row(line: str) -> list[str]:
    stripped = line.strip()
    if stripped.startswith("|"):
        stripped = stripped[1:]
    if stripped.endswith("|"):
        stripped = stripped[:-1]
    # Split on unescaped pipes.
    return [cell.strip().replace("\\|", "|") for cell in re.split(r"(?<!\\)\|", stripped)]


def _column_alignments(separator: str) -> list[str]:
    alignments = []
    for cell in _split_table_row(separator):
        left, right = cell.startswith(":"), cell.endswith(":")
        alignments.append("center" if left and right else "right" if right else "left")
    return alignments


def markdown_to_docx(markdown: str, output_path: Path, title: str | None = None) -> Path:
    """Write `markdown` to `output_path` as a .docx file."""
    writer = DocxWriter(title=title)
    render_markdown(writer, markdown)
    writer.save(output_path)
    return output_path


def render_markdown(writer: DocxWriter, markdown: str) -> None:  # noqa: C901 - line dispatcher
    lines = normalise_math_delimiters(markdown).replace("\r\n", "\n").split("\n")
    index = 0
    paragraph_buffer: list[str] = []
    quote_buffer: list[str] = []

    def flush_paragraph() -> None:
        nonlocal paragraph_buffer
        if paragraph_buffer:
            writer.add_paragraph_text(" ".join(part.strip() for part in paragraph_buffer))
            paragraph_buffer = []

    def flush_quote() -> None:
        nonlocal quote_buffer
        if quote_buffer:
            writer.add_quote(" ".join(part.strip() for part in quote_buffer))
            quote_buffer = []

    def flush_all() -> None:
        flush_paragraph()
        flush_quote()

    while index < len(lines):
        line = lines[index]
        stripped = line.strip()

        # Blank line — close the open block.
        if not stripped:
            flush_all()
            index += 1
            continue

        # Fenced code block.
        fence = FENCE_RE.match(line)
        if fence:
            flush_all()
            ticks = fence.group("ticks")[0]
            body: list[str] = []
            index += 1
            while index < len(lines):
                closing = FENCE_RE.match(lines[index])
                if closing and closing.group("ticks")[0] == ticks:
                    index += 1
                    break
                body.append(lines[index])
                index += 1
            writer.add_code_block(body)
            continue

        # Display maths: $$...$$ on one line, or spanning several lines.
        single_line = DISPLAY_MATH_RE.match(line)
        if single_line:
            flush_all()
            writer.add_display_math(single_line.group("body").strip())
            index += 1
            continue
        if stripped.startswith("$$"):
            flush_all()
            body = [stripped[2:]]
            index += 1
            while index < len(lines) and "$$" not in lines[index]:
                body.append(lines[index])
                index += 1
            if index < len(lines):
                body.append(lines[index].split("$$")[0])
                index += 1
            writer.add_display_math("\n".join(body).strip())
            continue

        # Picture, on a line of its own — a figure cut out of the page.
        picture = IMAGE_RE.match(line)
        if picture:
            flush_all()
            writer.add_picture(picture.group("path"), picture.group("alt").strip())
            index += 1
            continue

        # Heading.
        heading = HEADING_RE.match(line)
        if heading:
            flush_all()
            writer.add_heading(heading.group("text").strip(), len(heading.group("hashes")))
            index += 1
            continue

        # Horizontal rule.
        if HRULE_RE.match(line):
            flush_all()
            writer.add_rule()
            index += 1
            continue

        # Table.
        if TABLE_ROW_RE.match(line) and index + 1 < len(lines) and TABLE_SEP_RE.match(lines[index + 1]):
            flush_all()
            header = _split_table_row(line)
            alignments = _column_alignments(lines[index + 1])
            rows = [header]
            index += 2
            while index < len(lines) and TABLE_ROW_RE.match(lines[index]):
                rows.append(_split_table_row(lines[index]))
                index += 1
            writer.add_table(rows, alignments)
            continue

        # Blockquote.
        quote = QUOTE_RE.match(line)
        if quote:
            flush_paragraph()
            quote_buffer.append(quote.group("text"))
            index += 1
            continue

        # Lists.
        bullet = BULLET_RE.match(line)
        ordered = ORDERED_RE.match(line)
        if bullet or ordered:
            flush_all()
            match = bullet or ordered
            level = len(match.group("indent").replace("\t", "  ")) // 2
            text = match.group("text")
            # Absorb continuation lines that belong to this list item.
            index += 1
            while index < len(lines):
                follow = lines[index]
                if (
                    not follow.strip()
                    or BULLET_RE.match(follow)
                    or ORDERED_RE.match(follow)
                    or HEADING_RE.match(follow)
                    or follow.strip().startswith(("$$", "|", ">", "```"))
                ):
                    break
                text += " " + follow.strip()
                index += 1
            # The delimiter is not captured, so an ordered marker is put back
            # together; a bullet becomes the character Word would have drawn.
            marker = f"{match.group('marker')}." if ordered else "•"
            writer.add_list_item(text, marker=marker, ordered=ordered is not None, level=level)
            continue

        # Ordinary paragraph text.
        flush_quote()
        paragraph_buffer.append(line)
        index += 1

    flush_all()
