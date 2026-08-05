"""Write an analysed document as ordinary, editable Word content.

Everything here is real Word structure: paragraphs that reflow when a word is
added, headings Word's navigation pane can see, native equation objects, and
pictures in the text rather than floating over it. Nothing is positioned
absolutely, so the page will not be a facsimile of the original — a paragraph
that ran to four lines there may run to three here. What it will be is a document
someone can work in.

Formatting is carried across where the PDF states it: the typeface, size, weight,
slope and colour of every run are the ones the page was set in.
"""

from __future__ import annotations

from io import BytesIO
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import parse_xml
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

from .docx_builder import (
    ensure_styles,
    place_picture,
    render_inline,
    render_markdown,
    rule_border,
)
from .latex_omml import display_math_xml, inline_math_xml, looks_incomplete
from .pdf_extract import ImageItem, MathItem, Span, TableItem, TextLine
from .reflow import ImageBlock, MathBlock, Metrics, PageBlocks, TableBlock, TextBlock

# The least room a list number is given when the document does not say.
_MIN_LIST_INDENT_PT = 12.0
# However tightly the original was set, a printable margin is kept.
_MIN_MARGIN_PT = 18.0
# No block is set in further than this share of the column.
_MAX_INDENT_RATIO = 0.35
# Below this, a horizontal gap between two spans was a space the typesetter drew
# by moving the pen instead of by setting one.
_SPACE_RATIO = 0.2


def _base_size(lines: list[TextLine]) -> float:
    sizes = [span.size for line in lines for span in line.spans if span.text.strip()]
    return max(sizes) if sizes else 10.0


class StructuredWriter:
    """Builds a flowing .docx from analysed pages."""

    def __init__(self, metrics: Metrics, body_font: str = "Times New Roman") -> None:
        self.document = Document()
        ensure_styles(self.document)
        self.metrics = metrics
        normal = self.document.styles["Normal"]
        normal.font.name = body_font
        normal.font.size = Pt(metrics.size)
        # The template Word ships sets 1.15 line spacing. Over a page of a
        # document that was set solid it adds up to an extra inch and a half —
        # enough that the last figure of every page is pushed onto a page of its
        # own, which is then blank to the bottom.
        normal.paragraph_format.line_spacing = 1.0
        normal.paragraph_format.space_after = Pt(
            max(2.0, min(8.0, metrics.pitch - 1.2 * metrics.size))
        )
        normal.paragraph_format.space_before = Pt(0)
        # The page the document was set on, and the margins it was set within.
        # Left at the template's defaults the column is inches narrower than the
        # original, so every page overflows onto a second one and a picture that
        # fitted the source hangs off the edge of the copy.
        section = self.document.sections[0]
        section.page_width = Pt(metrics.page_width)
        section.page_height = Pt(metrics.page_height)
        section.left_margin = Pt(max(metrics.marker_left, _MIN_MARGIN_PT))
        section.right_margin = Pt(max(metrics.page_width - metrics.right, _MIN_MARGIN_PT))
        section.top_margin = Pt(metrics.top)
        section.bottom_margin = Pt(metrics.bottom)
        self.column = max(
            metrics.page_width - metrics.marker_left - (metrics.page_width - metrics.right),
            72.0,
        )
        self.list_indent = max(metrics.indent, _MIN_LIST_INDENT_PT)
        self._maths: list[MathItem] = []
        self._first_page = True
        self._break_pending = False
        body = self.document.element.body
        for paragraph in body.findall(qn("w:p")):
            body.remove(paragraph)

    # -- runs -------------------------------------------------------------- #

    def _style_run(self, run, span: Span, base: float = 0.0) -> None:
        run.font.name = span.font
        run.font.size = Pt(max(span.size, 1.0))
        run.bold = span.bold or None
        run.italic = span.italic or None
        # Here, unlike in the positioned writer, a raised span has to say so: the
        # text flows, so its coordinates are gone and only the mark-up is left to
        # lift it.
        if span.superscript:
            run.font.superscript = True
        elif span.subscript:
            run.font.subscript = True
        # Always stated, never inherited. A heading style carries a colour of its
        # own — Word's is blue — and a run that says nothing about its colour
        # takes it, so a black heading arrives blue.
        run.font.color.rgb = RGBColor.from_string(span.colour or "000000")
        fonts = run._r.get_or_add_rPr().find(qn("w:rFonts"))
        if fonts is not None:
            fonts.set(qn("w:cs"), span.font)

    def _add_math(self, paragraph, math_id: int, display: bool, align: str = "center") -> None:
        """Put an equation into a paragraph, natively if it can be read."""
        item = self._maths[math_id]
        if item.latex and not looks_incomplete(item.latex):
            try:
                xml = (
                    display_math_xml(item.latex, align) if display else inline_math_xml(item.latex)
                )
                paragraph._p.append(parse_xml(xml))
                return
            except Exception:
                pass
        # Unreadable: the crop is still exactly what the page showed.
        self._add_picture_run(paragraph, item.image, item.bbox)

    def _add_picture_run(self, paragraph, data: bytes, bbox) -> None:
        width = min(max(bbox[2] - bbox[0], 4.0), self.column)
        try:
            paragraph.add_run().add_picture(BytesIO(data), width=Pt(width))
        except Exception:
            pass  # an unreadable crop must not take the document down

    def _add_spans(self, paragraph, spans: list[Span], base: float) -> None:
        previous: Span | None = None
        for span in spans:
            if previous is not None and _needs_space(previous, span):
                paragraph.add_run(" ")
            if span.math_id is not None:
                self._add_math(paragraph, span.math_id, display=False)
            elif span.text:
                self._style_run(paragraph.add_run(span.text), span, base)
            previous = span

    def _add_lines(self, paragraph, lines: list[TextLine]) -> None:
        base = _base_size(lines)
        for index, line in enumerate(lines):
            if index:
                tail = _last_text(lines[index - 1])
                head = _first_text(line)
                if tail.endswith("-") and head[:1].islower():
                    # A word broken across the line break is put back together;
                    # in flowing text the break is gone, and the hyphen with it.
                    _strip_trailing_hyphen(paragraph)
                else:
                    paragraph.add_run(" ")
            self._add_spans(paragraph, line.spans, base)

    # -- blocks ------------------------------------------------------------ #

    def _paragraph(self, style: str | None = None):
        """Start a paragraph, carrying over a page break that is owed.

        The break is a property of the paragraph that follows it, not a
        paragraph of its own: an empty one carrying a break is still a line, and
        a page filled to its last line pushes that line — and the break with it
        — onto the next page, which then breaks again and comes out blank.
        """
        paragraph = self.document.add_paragraph(style=style)
        if self._break_pending:
            paragraph.paragraph_format.page_break_before = True
            self._break_pending = False
        return paragraph

    def _indent(self, block, centred: bool = False) -> float:
        """How far this block is set in from the left margin, as the page had it."""
        if centred:
            return 0.0
        offset = block.bbox[0] - self.metrics.marker_left
        if offset < 3.0:
            return 0.0
        return min(offset, _MAX_INDENT_RATIO * self.column)

    def _lead(self, paragraph, block, indent: float, span: Span | None) -> None:
        """Set the paragraph in, and write its list number in the margin it leaves.

        The original numbering is kept verbatim rather than handed to Word's
        automatic lists: in an answer key the numbers have to go on matching the
        questions, and a list Word renumbers would not.
        """
        fmt = paragraph.paragraph_format
        marker = getattr(block, "marker", "")
        hang = max(indent, self.list_indent) if marker else indent
        fmt.left_indent = Pt(hang)
        if not marker:
            return
        fmt.first_line_indent = Pt(-hang)
        run = paragraph.add_run(marker)
        if span is not None:
            self._style_run(run, span)
        paragraph.add_run("\t")

    def _text_block(self, block: TextBlock) -> None:
        head = block.lines[0].spans[0] if block.lines and block.lines[0].spans else None
        centred = block.align == "center"
        if block.kind == "heading":
            style = f"Heading {min(max(block.level, 1), 4)}"
            paragraph = self._paragraph(style if style in self.document.styles else None)
        else:
            paragraph = self._paragraph()
        if centred:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        self._lead(paragraph, block, self._indent(block, centred), head)
        self._add_lines(paragraph, block.lines)

    def _math_block(self, block: MathBlock) -> None:
        paragraph = self._paragraph()
        # Word sets an equation taller than the typesetter did, so the air
        # around it is kept to what tells it apart from the prose and no more.
        paragraph.paragraph_format.space_before = Pt(2)
        paragraph.paragraph_format.space_after = Pt(2)
        # A numbered equation is the answer to its question: it is set out from
        # its number, not centred away from it, and the number belongs in the
        # margin the list keeps for it.
        if block.marker:
            self._lead(paragraph, block, self._indent(block), None)
        else:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        self._add_math(
            paragraph, block.math_id, display=True, align="left" if block.marker else "center"
        )
        if block.label:
            paragraph.add_run("\t" + block.label)

    def _image_block(self, block: ImageBlock) -> None:
        paragraph = self._paragraph()
        if block.marker:
            self._lead(paragraph, block, self._indent(block), None)
        else:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        self._add_picture_run(paragraph, block.item.data, block.item.bbox)

    def _table_block(self, block: TableBlock) -> None:
        item: TableItem = block.item
        rows = [row for row in item.rows if row]
        if not rows:
            return
        columns = max(len(row) for row in rows)
        if self._break_pending:
            self._paragraph()  # a table cannot carry the break itself
        table = self.document.add_table(rows=0, cols=columns)
        table.style = "Table Grid" if "Table Grid" in self.document.styles else None
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        for index, cells in enumerate(rows):
            row = table.add_row()
            for column in range(columns):
                cell = row.cells[column]
                cell.text = ""
                paragraph = cell.paragraphs[0]
                paragraph.paragraph_format.space_after = Pt(0)
                value = cells[column] if column < len(cells) else ""
                if value:
                    run = paragraph.add_run(value.replace("\n", " "))
                    run.bold = True if index == 0 else None

    # -- Markdown pages ----------------------------------------------------- #
    #
    # A scanned page has no text layer to analyse: its words come back from the
    # vision model as Markdown, and the mark-up is the only structure the page
    # has. Rendering it through the same Markdown path the `flow` mode uses —
    # into this document, in its place in the page order — is what turns a scan
    # into headings, emphasis and native equations. Writing the Markdown out as
    # literal text instead is what made a scanned page arrive full of `**` and
    # `$\frac{}{}$`.
    #
    # The methods below are the interface `render_markdown` writes through; they
    # are the same set `DocxWriter` offers, implemented in this document's own
    # measurements so a transcribed page is set like the rest of it.

    def add_markdown_page(self, markdown: str) -> None:
        if markdown.strip():
            render_markdown(self, markdown)

    def add_paragraph_text(self, text: str, style: str | None = None):
        paragraph = self._paragraph(style if style and style in self.document.styles else None)
        render_inline(paragraph, text)
        return paragraph

    def add_heading(self, text: str, level: int) -> None:
        style = f"Heading {min(max(level, 1), 4)}"
        paragraph = self._paragraph(style if style in self.document.styles else None)
        render_inline(paragraph, text)

    def add_rule(self) -> None:
        rule_border(self._paragraph())

    def add_quote(self, text: str) -> None:
        style = "Figure Note" if text.lstrip().startswith(("_Figure", "*Figure")) else "Quote"
        if style not in self.document.styles:
            style = "Figure Note"
        self.add_paragraph_text(text, style=style)

    def add_list_item(self, text: str, marker: str, ordered: bool, level: int) -> None:
        """Set a list item out from its own marker, kept verbatim.

        Word's automatic numbering is deliberately not used, for the reason the
        extracted path avoids it too: in an answer key the numbers have to go on
        matching the questions, and a list Word renumbers would not.
        """
        paragraph = self._paragraph()
        hang = self.list_indent * (level + 1)
        paragraph.paragraph_format.left_indent = Pt(hang)
        paragraph.paragraph_format.first_line_indent = Pt(-self.list_indent)
        paragraph.add_run(marker)
        paragraph.add_run("\t")
        render_inline(paragraph, text)

    def add_display_math(self, latex: str) -> None:
        paragraph = self._paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.space_before = Pt(2)
        paragraph.paragraph_format.space_after = Pt(2)
        try:
            paragraph._p.append(parse_xml(display_math_xml(latex)))
        except Exception:
            # Unreadable LaTeX still says more as itself than as an empty slot.
            paragraph.add_run(latex)

    def add_code_block(self, lines: list[str]) -> None:
        for line in lines or [""]:
            paragraph = self._paragraph("Code Block" if "Code Block" in self.document.styles else None)
            paragraph.add_run(line)

    def add_table(self, rows: list[list[str]], alignments: list[str]) -> None:
        if not rows:
            return
        columns = max(len(row) for row in rows)
        if self._break_pending:
            self._paragraph()  # a table cannot carry the break itself
        table = self.document.add_table(rows=0, cols=columns)
        table.style = "Table Grid" if "Table Grid" in self.document.styles else None
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        for index, cells in enumerate(rows):
            row = table.add_row()
            for column in range(columns):
                cell = row.cells[column]
                cell.text = ""
                paragraph = cell.paragraphs[0]
                paragraph.paragraph_format.space_after = Pt(0)
                if column < len(alignments):
                    paragraph.alignment = {
                        "center": WD_ALIGN_PARAGRAPH.CENTER,
                        "right": WD_ALIGN_PARAGRAPH.RIGHT,
                    }.get(alignments[column], WD_ALIGN_PARAGRAPH.LEFT)
                value = cells[column] if column < len(cells) else ""
                if value:
                    render_inline(paragraph, value, bold=(index == 0))

    def add_picture(self, path: str, caption: str = "") -> None:
        paragraph = self._paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if not place_picture(paragraph, path, self.column, self.metrics.page_height * 0.6):
            return
        if caption:
            self.add_quote(f"_{caption}_")

    def add_page_break(self) -> None:
        self._break_pending = True

    # -- pages -------------------------------------------------------------- #

    def add_page(self, page: PageBlocks, maths: list[MathItem]) -> None:
        self._maths = maths
        # The text flows on from where the page before it ended. Breaking at
        # every page boundary instead only looks right while every page happens
        # to fit: Word sets an equation taller than the typesetter did, so a
        # source page that was full here runs a line or two over — and that line
        # is then followed by a forced break, leaving nine tenths of the page it
        # landed on empty. Where the source really does start something new, the
        # heading it starts with is what says so.
        if not page.blocks:
            # A blank page is a break the typesetter put in deliberately; it is
            # the one thing on it worth keeping.
            self._break_pending = self._break_pending or (page.blank and not self._first_page)
            return
        if not self._first_page and _starts_section(page):
            self._break_pending = True
        self._first_page = False
        for block in page.blocks:
            if isinstance(block, TextBlock):
                self._text_block(block)
            elif isinstance(block, MathBlock):
                self._math_block(block)
            elif isinstance(block, ImageBlock):
                self._image_block(block)
            elif isinstance(block, TableBlock):
                self._table_block(block)

    def save(self, path: Path) -> None:
        self.document.save(path)


# --------------------------------------------------------------------------- #
# Helpers
# --------------------------------------------------------------------------- #


def _starts_section(page: PageBlocks) -> bool:
    """Does this page open with a heading, rather than carry on the last one?

    Read past the pictures: a section that opens with a masthead across the head
    of the page still opens with the heading printed under it.
    """
    for block in page.blocks:
        if isinstance(block, TextBlock):
            return block.kind == "heading"
    return False


def _needs_space(previous: Span, following: Span) -> bool:
    if previous.text.endswith(" ") or following.text.startswith(" "):
        return False
    if not previous.text or not following.text:
        return False
    gap = following.bbox[0] - previous.bbox[2]
    return gap > _SPACE_RATIO * max(previous.size, 1.0)


def _first_text(line: TextLine) -> str:
    for span in line.spans:
        if span.text.strip():
            return span.text.lstrip()
    return ""


def _last_text(line: TextLine) -> str:
    for span in reversed(line.spans):
        if span.text.strip():
            return span.text.rstrip()
    return ""


def _strip_trailing_hyphen(paragraph) -> None:
    for run in reversed(paragraph.runs):
        if run.text.rstrip().endswith("-"):
            run.text = run.text.rstrip()[:-1]
            return
        if run.text.strip():
            return


def write_document(
    pages: list[tuple[PageBlocks, list[MathItem]]],
    metrics: Metrics,
    path: Path,
    body_font: str = "Times New Roman",
) -> Path:
    writer = StructuredWriter(metrics, body_font=body_font)
    for page, maths in pages:
        writer.add_page(page, maths)
    writer.save(path)
    return path


def dominant_font(pages) -> str:
    """The typeface most of the document's text is set in."""
    from collections import Counter

    counts: Counter[str] = Counter()
    for page in pages:
        for block in page.blocks:
            if isinstance(block, TextBlock):
                for line in block.lines:
                    for span in line.spans:
                        counts[span.font] += len(span.text)
    return counts.most_common(1)[0][0] if counts else "Times New Roman"
