"""Write a positioned layout model into a .docx that mirrors the original page.

Word is a flow format, so the only way to reproduce a fixed page faithfully is to
stop flowing: every piece of the page becomes a floating shape anchored to an
absolute position on the sheet. Each PDF page gets its own section sized exactly
to the source page with zero margins, and text, pictures, rules, tables and
equations are placed at the coordinates they had in the PDF.

The trade-off this buys — and it is the point of the mode — is that editing a
block does not reflow its neighbours.
"""

from __future__ import annotations

from io import BytesIO
from pathlib import Path

import fitz
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsmap, qn
from docx.shared import Pt, RGBColor

from .latex_omml import display_math_xml, inline_math_xml, looks_incomplete
from .pdf_extract import ImageItem, MathItem, PageLayout, RuleItem, TableItem, TextLine

# python-docx only knows the OOXML prefixes; VML positioning needs these too, and
# registering them here is what lets `qn()` resolve v:/o:/w10: names below.
nsmap.setdefault("v", "urn:schemas-microsoft-com:vml")
nsmap.setdefault("o", "urn:schemas-microsoft-com:office:office")
nsmap.setdefault("w10", "urn:schemas-microsoft-com:office:word")

# Namespaces every hand-built shape needs to declare.
_NS = (
    'xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" '
    'xmlns:v="urn:schemas-microsoft-com:vml" '
    'xmlns:o="urn:schemas-microsoft-com:office:office" '
    'xmlns:w10="urn:schemas-microsoft-com:office:word" '
    'xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships"'
)

# Stacking order. Word paints higher z-index later; a negative index sits behind
# the text layer, which is what puts an OCR layer under a scanned page image.
Z_SCAN_TEXT = -1
Z_IMAGE = 2
Z_RULE = 3
Z_TABLE = 4
Z_MATH = 5
Z_TEXT = 6

# Word refuses page dimensions beyond 22in; clamp rather than emit an invalid file.
_MAX_PAGE_PT = 1584.0
_MIN_RULE_PT = 0.5

# Base-14 metric stand-ins used to predict how wide Word will draw a run.
_METRIC_FONTS = {
    "Times New Roman": ("tiro", "tibo", "tiit", "tibi"),
    "Arial": ("helv", "hebo", "heit", "hebi"),
    "Courier New": ("cour", "cobo", "coit", "cobi"),
}
# Beyond this the measurement is not to be trusted, and stretching would look
# worse than the drift it is correcting.
_MAX_TRACKING_RATIO = 0.4

# How far a frame's text may be predicted to wander before it is broken into a
# second frame that starts from an exact coordinate again. Character spacing
# corrects most of the difference between the real font and the substitute, but
# not identically in every reader — Word adds it after every character, others
# only between them — so the residue is capped rather than trusted.
_DRIFT_BUDGET_PT = 3.0


def _metric_font(font: str, bold: bool, italic: bool) -> str:
    faces = _METRIC_FONTS.get(font) or _METRIC_FONTS["Times New Roman"]
    return faces[(1 if bold else 0) + (2 if italic else 0)]


def _measure(text: str, font: str, size: float, bold: bool, italic: bool) -> float:
    """Width in points that Word will give `text`, via base-14 metrics."""
    try:
        return fitz.get_text_length(text, fontname=_metric_font(font, bold, italic), fontsize=size)
    except Exception:
        return 0.0


def _pt(value: float) -> str:
    return f"{round(value, 2)}pt"


def _style(x: float, y: float, width: float, height: float, z: int) -> str:
    return (
        "position:absolute;"
        f"margin-left:{_pt(x)};margin-top:{_pt(y)};"
        f"width:{_pt(max(width, 0.5))};height:{_pt(max(height, 0.5))};"
        f"z-index:{z};"
        "mso-position-horizontal-relative:page;"
        "mso-position-vertical-relative:page;"
        "mso-wrap-style:none;v-text-anchor:top"
    )


class ReplicaWriter:
    """Builds a page-for-page visual replica of a PDF."""

    def __init__(self) -> None:
        self.document = Document()
        self._shape_id = 0
        self._page_index = 0
        self._page_width = _MAX_PAGE_PT
        # A fresh document may carry an empty paragraph; drop it so page one is
        # not pushed down by a stray line.
        body = self.document.element.body
        for paragraph in body.findall(qn("w:p")):
            body.remove(paragraph)

    # -- scaffolding ------------------------------------------------------- #

    def _next_id(self) -> int:
        self._shape_id += 1
        return self._shape_id

    def _section_for(self, width: float, height: float):
        if self._page_index == 0:
            section = self.document.sections[0]
        else:
            section = self.document.add_section(WD_SECTION.NEW_PAGE)
        section.page_width = Pt(min(width, _MAX_PAGE_PT))
        section.page_height = Pt(min(height, _MAX_PAGE_PT))
        for margin in ("left_margin", "right_margin", "top_margin", "bottom_margin"):
            setattr(section, margin, Pt(0))
        section.header_distance = Pt(0)
        section.footer_distance = Pt(0)
        self._page_index += 1
        return section

    def _anchor(self):
        """A near-invisible paragraph that all of a page's shapes hang from."""
        paragraph = self.document.add_paragraph()
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.line_spacing = Pt(1)
        return paragraph

    def _detach(self, element):
        """Take an element back out of the body so it can be re-homed in a shape."""
        parent = element.getparent()
        if parent is not None:
            parent.remove(element)
        return element

    # -- shape primitives -------------------------------------------------- #

    def _textbox(self, anchor, box, z: int, blocks: list):
        """Place `blocks` (w:p / w:tbl elements) in a frame at `box`."""
        x0, y0, x1, y1 = box
        shape = parse_xml(
            f"<w:pict {_NS}>"
            f'<v:rect id="s{self._next_id()}" style="{_style(x0, y0, x1 - x0, y1 - y0, z)}" '
            'filled="f" stroked="f">'
            '<v:textbox inset="0,0,0,0" style="mso-fit-shape-to-text:f">'
            "<w:txbxContent/></v:textbox>"
            '<w10:wrap type="none" anchorx="page" anchory="page"/>'
            "</v:rect></w:pict>"
        )
        content = shape.find(qn("v:rect")).find(qn("v:textbox")).find(qn("w:txbxContent"))
        for block in blocks:
            content.append(self._detach(block))

        run = anchor.add_run()
        run.font.size = Pt(1)
        run._r.append(shape)

    def _rule(self, anchor, rule: RuleItem):
        x0, y0, x1, y1 = rule.bbox
        width = max(x1 - x0, _MIN_RULE_PT)
        height = max(y1 - y0, _MIN_RULE_PT)
        shape = parse_xml(
            f"<w:pict {_NS}>"
            f'<v:rect id="s{self._next_id()}" style="{_style(x0, y0, width, height, Z_RULE)}" '
            f'fillcolor="#{rule.colour}" stroked="f">'
            '<w10:wrap type="none" anchorx="page" anchory="page"/>'
            "</v:rect></w:pict>"
        )
        run = anchor.add_run()
        run.font.size = Pt(1)
        run._r.append(shape)

    def _tight_paragraph(self, height: float | None = None):
        """A paragraph with every scrap of default spacing removed."""
        paragraph = self.document.add_paragraph()
        fmt = paragraph.paragraph_format
        fmt.space_before = Pt(0)
        fmt.space_after = Pt(0)
        if height:
            # Exact leading keeps the baseline where the PDF had it.
            fmt.line_spacing_rule = None
            spacing = paragraph._p.get_or_add_pPr().get_or_add_spacing()
            spacing.set(qn("w:line"), str(int(round(height * 20))))
            spacing.set(qn("w:lineRule"), "exact")
        return paragraph

    # -- content ----------------------------------------------------------- #

    def _picture_shape(self, anchor, box, data: bytes, z: int) -> bool:
        """Float a picture directly, with no text frame around it.

        Wrapping a picture in a text box costs it the frame's internal padding —
        a few points of drift on every image, diagram and equation — so the image
        part is related to the document and referenced by a VML shape instead.
        """
        x0, y0, x1, y1 = box
        try:
            rid, _ = self.document.part.get_or_add_image(BytesIO(data))
        except Exception:
            return False
        shape = parse_xml(
            f"<w:pict {_NS}>"
            f'<v:rect id="p{self._next_id()}" '
            f'style="{_style(x0, y0, x1 - x0, y1 - y0, z)}" stroked="f">'
            f'<v:imagedata r:id="{rid}" o:title=""/>'
            '<w10:wrap type="none" anchorx="page" anchory="page"/>'
            "</v:rect></w:pict>"
        )
        run = anchor.add_run()
        run.font.size = Pt(1)
        run._r.append(shape)
        return True

    @staticmethod
    def _relaxed(box, pad: float = 6.0):
        """Grow a frame down and right so nothing is clipped by its own border."""
        x0, y0, x1, y1 = box
        return (x0, y0, x1 + pad, y1 + pad)

    def add_image(self, anchor, item: ImageItem):
        # An unreadable image must not take the whole page down.
        self._picture_shape(anchor, item.bbox, item.data, Z_IMAGE)

    @staticmethod
    def _track(run, span, advance: float | None) -> None:
        """Stretch or tighten a run so it occupies the width the PDF gave it.

        The reader almost never has the document's original font, and a
        substitute of different width would walk every later word on the line out
        of place. Word's character spacing pins the run back to its true advance.
        """
        if not advance or not span.text:
            return
        natural = _measure(span.text, span.font, span.size, span.bold, span.italic)
        if natural <= 0:
            return
        per_character = (advance - natural) / len(span.text)
        if abs(per_character) > _MAX_TRACKING_RATIO * max(span.size, 1.0):
            return
        twentieths = int(round(per_character * 20))
        if twentieths == 0:
            return

        spacing = parse_xml(
            '<w:spacing xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" '
            f'w:val="{twentieths}"/>'
        )
        # w:spacing has a fixed slot in the run-properties sequence — after
        # w:color, before w:sz. Out of order, Word rejects the part outright.
        rpr = run._r.get_or_add_rPr()
        anchor = rpr.find(qn("w:sz"))
        if anchor is not None:
            anchor.addprevious(spacing)
            return
        anchor = rpr.find(qn("w:color"))
        if anchor is not None:
            anchor.addnext(spacing)
            return
        rpr.append(spacing)

    def _add_runs(self, paragraph, spans, track: bool = False) -> None:
        for index, span in enumerate(spans):
            if not span.text:
                continue
            run = paragraph.add_run(span.text)
            run.font.name = span.font
            run.bold = span.bold or None
            run.italic = span.italic or None
            run.font.color.rgb = RGBColor.from_string(span.colour)
            # Superscript is deliberately not applied. A raised span already
            # carries its own reduced size and its own raised coordinates, and
            # this writer honours both — asking Word to superscript it as well
            # shrinks and lifts it a second time. The flag is unreliable besides:
            # the extractor sets it whenever a span sits above the line's
            # dominant baseline, so an ordinary sentence sharing a line with the
            # low half of a fraction is reported as superscript throughout.
            run.font.size = Pt(max(span.size, 1.0))

            # Word resolves the East-Asian and complex-script faces separately.
            fonts = run._r.get_or_add_rPr().find(qn("w:rFonts"))
            if fonts is not None:
                fonts.set(qn("w:cs"), span.font)
                fonts.set(qn("w:eastAsia"), span.font)

            if track:
                # Match each run's advance to where the next run starts, so the
                # gaps between words are reproduced along with the words.
                following = spans[index + 1] if index + 1 < len(spans) else None
                advance = (
                    following.bbox[0] - span.bbox[0]
                    if following
                    else span.bbox[2] - span.bbox[0]
                )
                self._track(run, span, advance)

    @staticmethod
    def _split_drift(spans: list) -> list[list]:
        """Cut a run of spans wherever predicted drift would become visible.

        Each resulting frame is anchored at a true PDF coordinate, so error can
        never accumulate across a whole line the way it would in one long frame.
        """
        parts: list[list] = []
        current: list = []
        budget = 0.0
        for index, span in enumerate(spans):
            following = spans[index + 1] if index + 1 < len(spans) else None
            advance = (
                following.bbox[0] - span.bbox[0] if following else span.bbox[2] - span.bbox[0]
            )
            error = abs(advance - _measure(span.text, span.font, span.size, span.bold, span.italic))
            if current and budget + error > _DRIFT_BUDGET_PT:
                parts.append(current)
                current, budget = [], 0.0
            current.append(span)
            budget += error
        if current:
            parts.append(current)
        return parts

    def add_line(self, anchor, line: TextLine) -> None:
        """Place a line as one frame per stretch of ordinary text.

        Equations are laid out separately at their own coordinates, so a line is
        split around them. Keeping each stretch in its own frame means the words
        after an equation stay exactly where the PDF put them.
        """
        segments: list[list] = []
        current: list = []
        for span in line.spans:
            if span.math_id is not None:
                if current:
                    segments.append(current)
                    current = []
                continue
            current.append(span)
        if current:
            segments.append(current)

        for spans in [part for segment in segments for part in self._split_drift(segment)]:
            if not any(span.text.strip() for span in spans):
                continue
            x0 = min(span.bbox[0] for span in spans)
            y0 = min(span.bbox[1] for span in spans)
            x1 = max(span.bbox[2] for span in spans)
            y1 = max(span.bbox[3] for span in spans)
            height = max(y1 - y0, 1.0)
            paragraph = self._tight_paragraph(height)
            self._add_runs(paragraph, spans, track=True)
            # Word will not have the PDF's exact font, so the same string can
            # measure wider. Frames get generous slack — they are invisible — so
            # text never wraps to a second line and collides with the line below.
            width = max((x1 - x0) * 1.35, (x1 - x0) + 24.0)
            # That slack must not push the frame off the sheet. A frame hanging
            # over the edge is inside its rights in Word but makes other readers
            # shrink the text to fit the page, and a line near the right margin
            # comes out at half size. Give the slack back rather than the words.
            width = max(min(width, self._page_width - x0), (x1 - x0) + 2.0)
            # Anchored on the line's own top edge — measured against the source,
            # any padding here shows up directly as vertical drift.
            self._textbox(anchor, (x0, y0, x0 + width, y1 + 2.0), Z_TEXT, [paragraph._p])

    def add_math(self, anchor, item: MathItem) -> None:
        x0, y0, x1, y1 = item.bbox
        # A transcription that stops mid-expression is drawn by Word as an empty
        # slot; the crop it came from is still exactly what the page showed.
        if item.latex and not looks_incomplete(item.latex):
            try:
                paragraph = self._tight_paragraph()
                xml = display_math_xml(item.latex) if item.display else inline_math_xml(item.latex)
                paragraph._p.append(parse_xml(xml))
                self._textbox(anchor, self._relaxed((x0, y0 - 1.0, x1, y1), 10.0), Z_MATH, [paragraph._p])
                return
            except Exception:
                pass
        self._picture_shape(anchor, item.bbox, item.image, Z_MATH)

    def add_table(self, anchor, item: TableItem) -> None:
        rows = [row for row in item.rows if row]
        if not rows:
            return
        columns = max(len(row) for row in rows)
        table = self.document.add_table(rows=0, cols=columns)
        table.style = "Table Grid" if "Table Grid" in self.document.styles else "Normal Table"
        table.alignment = WD_TABLE_ALIGNMENT.LEFT
        table.autofit = False

        x0, y0, x1, y1 = item.bbox
        widths = self._column_widths(item, columns, x1 - x0)
        for cells in rows:
            row = table.add_row()
            for index in range(columns):
                cell = row.cells[index]
                cell.width = Pt(widths[index])
                cell.text = ""
                paragraph = cell.paragraphs[0]
                paragraph.paragraph_format.space_before = Pt(0)
                paragraph.paragraph_format.space_after = Pt(0)
                value = cells[index] if index < len(cells) else ""
                if value:
                    paragraph.add_run(value.replace("\n", " "))

        block = table._tbl
        # Word wants a paragraph after a table inside a text frame.
        tail = self._tight_paragraph()
        self._textbox(anchor, (x0, y0, x1, y1), Z_TABLE, [block, tail._p])

    @staticmethod
    def _column_widths(item: TableItem, columns: int, total: float) -> list[float]:
        edges = sorted({round(x, 1) for x in item.col_x})
        if len(edges) == columns + 1:
            widths = [edges[i + 1] - edges[i] for i in range(columns)]
            if all(width > 0 for width in widths):
                return widths
        share = max(total / columns, 1.0) if columns else total
        return [share] * columns

    # -- page -------------------------------------------------------------- #

    def add_page(self, layout: PageLayout) -> None:
        self._page_width = min(layout.width, _MAX_PAGE_PT)
        self._section_for(layout.width, layout.height)
        anchor = self._anchor()

        if layout.scanned:
            # The raster is the page. Any transcribed text goes behind it, so the
            # page looks exactly like the original but the words are still there.
            if layout.lines:
                blocks = []
                for line in layout.lines:
                    paragraph = self._tight_paragraph()
                    self._add_runs(paragraph, line.spans)
                    blocks.append(paragraph._p)
                self._textbox(
                    anchor, (0, 0, layout.width, layout.height), Z_SCAN_TEXT, blocks
                )
            if layout.page_image:
                self.add_image(
                    anchor, ImageItem(bbox=(0, 0, layout.width, layout.height), data=layout.page_image)
                )
            return

        for image in layout.images:
            self.add_image(anchor, image)
        for rule in layout.rules:
            self._rule(anchor, rule)
        for table in layout.tables:
            self.add_table(anchor, table)

        # Every equation — display or inline — is placed at its own coordinates.
        for item in layout.maths:
            if item.dropped:
                continue  # not an equation after all; its text is drawn below
            self.add_math(anchor, item)
        for line in layout.lines:
            if line.math_id is not None:
                continue  # the whole line is an equation, already placed
            self.add_line(anchor, line)

    def save(self, path: Path) -> None:
        self.document.save(path)
