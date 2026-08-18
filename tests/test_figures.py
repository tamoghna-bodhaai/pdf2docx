"""A figure must come out of the page at the size, and the detail, it went in at.

Two things went wrong here and they compounded. The crop was rendered at three
times the resolution the page could supply, and then written to a file that
understated its resolution, so Word placed it three times too large as well. A
figure an inch and a half across arrived four and a half inches across, carrying
no more detail than the inch and a half had.

Both are arithmetic, so both are testable without a model in the loop: the
fixtures are pages of known resolution and the assertions are about the pixels
and the header of the file that comes out.
"""

from __future__ import annotations

from dataclasses import replace
from pathlib import Path

import fitz
import pytest
from docx import Document
from docx.shared import Emu

from app import figures
from app.config import settings
from app.docx_builder import DocxWriter, place_picture, render_markdown, resolve_picture

# Somewhere in the middle of the page, well clear of every edge.
BOX = "200,200,400,400"
MARKDOWN = f"> _Figure: a diagram of something_ <!--box: {BOX}-->"

DIAGRAM_DPI = 300
SCAN_DPI = 90.0


@pytest.fixture(autouse=True)
def configured(monkeypatch):
    """Pin the settings the crop depends on; `Settings` is frozen, so replace it."""

    def configure(**overrides):
        pinned = {"diagram_dpi": DIAGRAM_DPI, "crop_native_dpi": True, **overrides}
        monkeypatch.setattr(figures, "settings", replace(settings, **pinned))

    configure()
    return configure


def reference(pdf: Path, work_dir: Path) -> str:
    """Run the figure placer over one page and return the reference it wrote."""
    result = figures.place_figures(pdf, 0, MARKDOWN, work_dir)
    assert result.startswith("!["), f"no crop was made: {result!r}"
    return result[result.index("(") + 1 : -1]


def crop(pdf: Path, work_dir: Path) -> Path:
    """The crop file itself, resolved the way the document writers resolve it."""
    path = resolve_picture(work_dir, reference(pdf, work_dir))
    assert path is not None and path.exists()
    return path


def rect_of(pdf: Path) -> fitz.Rect:
    """The page rectangle `BOX` resolves to — the size the crop should end up."""
    with fitz.open(pdf) as document:
        page = document.load_page(0)
        rect = figures._rect(BOX, page.rect.width, page.rect.height)
    assert rect is not None
    return rect


def natural_size_pt(png: Path) -> tuple[float, float]:
    """How large Word will draw the file if nobody tells it otherwise."""
    document = Document()
    picture = document.add_paragraph().add_run().add_picture(str(png))
    return Emu(picture.width).pt, Emu(picture.height).pt


# -- the sizing bug --------------------------------------------------------- #


def test_crop_is_placed_at_the_size_the_page_drew_it(scan, tmp_path):
    """The regression: a crop used to arrive `diagram_dpi / 96` times too large."""
    pdf = scan(dpi=SCAN_DPI)
    png = crop(pdf, tmp_path)
    rect = rect_of(pdf)

    width, height = natural_size_pt(png)
    assert width == pytest.approx(rect.width, rel=0.02)
    assert height == pytest.approx(rect.height, rel=0.02)


def test_crop_declares_the_resolution_it_was_rendered_at(scan, tmp_path):
    """Pixels over declared resolution is the whole of Word's size arithmetic."""
    pdf = scan(dpi=SCAN_DPI)
    png = crop(pdf, tmp_path)
    rect = rect_of(pdf)

    pixmap = fitz.Pixmap(str(png))
    assert pixmap.xres == pixmap.yres
    assert pixmap.width / pixmap.xres * 72.0 == pytest.approx(rect.width, rel=0.02)


def test_a_drawn_figure_is_also_placed_at_its_true_size(drawn, tmp_path):
    """The vector path has no resolution ceiling, and must still size correctly."""
    png = crop(drawn, tmp_path)
    rect = rect_of(drawn)

    width, _ = natural_size_pt(png)
    assert width == pytest.approx(rect.width, rel=0.02)


def test_a_figure_wider_than_the_column_is_still_shrunk(scan, tmp_path):
    """Sizing honestly must not cost the fit-to-column guard its job."""
    png = crop(scan(dpi=SCAN_DPI), tmp_path)
    column = 40.0  # narrower than any crop this test makes

    document = Document()
    assert place_picture(document.add_paragraph(), str(png), column)
    assert Emu(document.inline_shapes[0].width).pt <= column + 0.5


# -- the resolution bug ----------------------------------------------------- #


def test_a_scan_is_not_rendered_beyond_its_own_resolution(scan, tmp_path):
    """Cutting a 90 DPI scan at 300 DPI interpolates; it does not recover detail."""
    pdf = scan(dpi=SCAN_DPI)
    png = crop(pdf, tmp_path)
    rect = rect_of(pdf)

    pixmap = fitz.Pixmap(str(png))
    assert pixmap.width == pytest.approx(rect.width * SCAN_DPI / 72.0, abs=2)
    assert pixmap.width < rect.width * DIAGRAM_DPI / 72.0 / 2


def test_a_drawn_figure_keeps_full_diagram_resolution(drawn, tmp_path):
    """Vector artwork can be rendered as large as anyone likes and stays sharp."""
    png = crop(drawn, tmp_path)
    rect = rect_of(drawn)

    pixmap = fitz.Pixmap(str(png))
    assert pixmap.width == pytest.approx(rect.width * DIAGRAM_DPI / 72.0, abs=2)


def test_text_over_a_scan_lifts_the_ceiling(scan, tmp_path):
    """A region carrying real text is not purely a scan, so it is cut at full DPI."""
    pdf = scan(dpi=SCAN_DPI, name="scan-with-text.pdf")
    with fitz.open(pdf) as document:
        page = document.load_page(0)
        page.insert_text(fitz.Point(150, 250), "a real text layer", fontsize=11)
        document.save(str(pdf), incremental=True, encryption=fitz.PDF_ENCRYPT_KEEP)

    png = crop(pdf, tmp_path)
    pixmap = fitz.Pixmap(str(png))
    assert pixmap.width == pytest.approx(rect_of(pdf).width * DIAGRAM_DPI / 72.0, abs=2)


def test_a_very_coarse_scan_is_floored_not_honoured(scan, tmp_path):
    """Honouring 20 DPI would hand Word an image too small to draw."""
    pdf = scan(dpi=20.0, name="coarse.pdf")
    png = crop(pdf, tmp_path)
    rect = rect_of(pdf)

    pixmap = fitz.Pixmap(str(png))
    assert pixmap.width == pytest.approx(rect.width * figures._MIN_CROP_DPI / 72.0, abs=2)


def test_a_page_that_will_not_parse_still_yields_a_figure(scan, tmp_path, monkeypatch):
    """Measuring the page is an optimisation; failing to must not cost the crop."""

    def unparseable(self, *args, **kwargs):
        raise RuntimeError("cannot parse the content stream")

    monkeypatch.setattr(fitz.Page, "get_drawings", unparseable)

    pdf = scan(dpi=SCAN_DPI)
    png = crop(pdf, tmp_path)
    rect = rect_of(pdf)

    # No ceiling could be established, so the configured resolution stands —
    # but the crop must still declare the size it really is.
    pixmap = fitz.Pixmap(str(png))
    assert pixmap.width == pytest.approx(rect.width * DIAGRAM_DPI / 72.0, abs=2)
    assert natural_size_pt(png)[0] == pytest.approx(rect.width, rel=0.02)


def test_the_cap_can_be_switched_off(scan, tmp_path, configured):
    """`PDF2DOCX_CROP_NATIVE=off` restores cutting everything at `diagram_dpi`."""
    configured(crop_native_dpi=False)
    pdf = scan(dpi=SCAN_DPI)
    png = crop(pdf, tmp_path)

    pixmap = fitz.Pixmap(str(png))
    assert pixmap.width == pytest.approx(rect_of(pdf).width * DIAGRAM_DPI / 72.0, abs=2)


def test_switching_the_cap_off_does_not_bring_back_the_sizing_bug(scan, tmp_path, configured):
    """The two fixes are independent: resolution is a choice, size is a fact."""
    configured(crop_native_dpi=False)
    pdf = scan(dpi=SCAN_DPI)
    png = crop(pdf, tmp_path)

    width, _ = natural_size_pt(png)
    assert width == pytest.approx(rect_of(pdf).width, rel=0.02)


# -- the guards that were already there ------------------------------------- #


@pytest.mark.parametrize(
    "box",
    [
        "400,400,200,200",  # inside out
        "500,500,505,505",  # a sliver
        "0,0,1000,1000",  # the whole page
        "1,2,3",  # not four numbers
        "a,b,c,d",  # not numbers
    ],
)
def test_an_implausible_box_keeps_its_caption(scan, tmp_path, box):
    result = figures.place_figures(
        scan(dpi=SCAN_DPI), 0, f"> _Figure: a diagram_ <!--box: {box}-->", tmp_path
    )
    assert "<!--" not in result
    assert "![" not in result
    assert "a diagram" in result


def test_a_page_past_the_end_is_left_alone(scan, tmp_path):
    result = figures.place_figures(scan(dpi=SCAN_DPI), 99, MARKDOWN, tmp_path)
    assert "<!--" not in result
    assert "![" not in result


# -- the reference the crop leaves behind ----------------------------------- #
#
# The Markdown outlives the process that wrote it: it is saved beside the
# document, handed back over the API and downloadable. A reference into this
# machine's filesystem is both useless anywhere else and a description of this
# machine, so the reference is relative and the reader resolves it.


def test_the_reference_is_relative_to_the_working_directory(scan, tmp_path):
    """The regression: crops used to be referred to by their absolute path."""
    link = reference(scan(dpi=SCAN_DPI), tmp_path)

    assert link == f"{figures.FIGURE_DIR}/page-0001-figure-1.png"
    assert not Path(link).is_absolute()
    assert str(tmp_path) not in link


def test_the_crop_is_written_under_the_working_directory(scan, tmp_path):
    """Relative to what, specifically: the directory the caller handed over."""
    link = reference(scan(dpi=SCAN_DPI), tmp_path)
    assert (tmp_path / link).exists()


def test_the_markdown_names_no_host_path(scan, tmp_path):
    """What a downloaded `.md` is allowed to say about the machine that made it."""
    result = figures.place_figures(scan(dpi=SCAN_DPI), 0, MARKDOWN, tmp_path)

    assert str(tmp_path) not in result
    assert str(Path.home()) not in result


def test_a_relative_reference_resolves_under_the_working_directory(tmp_path):
    figure = tmp_path / figures.FIGURE_DIR / "page-0001-figure-1.png"
    figure.parent.mkdir(parents=True)
    figure.write_bytes(b"")

    assert resolve_picture(tmp_path, f"{figures.FIGURE_DIR}/page-0001-figure-1.png") == figure


@pytest.mark.parametrize(
    "src",
    [
        "../../../etc/passwd",  # climbing out
        "/etc/passwd",  # absolute, which `/` joining silently honours
        "figures/../../secret.pdf",  # out and back down
        ".",  # the working directory itself
    ],
)
def test_a_reference_outside_the_working_directory_is_refused(tmp_path, src):
    """The Markdown came from a model reading an untrusted page."""
    assert resolve_picture(tmp_path, src) is None


def test_without_a_working_directory_a_reference_is_taken_as_it_stands(tmp_path):
    """Nothing to resolve against and nothing to confine to — library use."""
    assert resolve_picture(None, "/somewhere/else.png") == Path("/somewhere/else.png")


# -- the two ends joined up -------------------------------------------------- #


def test_a_relative_reference_survives_the_round_trip_into_word(scan, tmp_path):
    """Cropped here, resolved there: the whole point of the reference being relative."""
    pdf = scan(dpi=SCAN_DPI)
    markdown = figures.place_figures(pdf, 0, MARKDOWN, tmp_path)

    writer = DocxWriter(base_dir=tmp_path)
    render_markdown(writer, markdown)
    writer.save(tmp_path / "out.docx")

    document = Document(str(tmp_path / "out.docx"))
    assert len(document.inline_shapes) == 1
    assert Emu(document.inline_shapes[0].width).pt == pytest.approx(rect_of(pdf).width, rel=0.02)


def test_a_picture_after_an_empty_numbered_item_is_not_swallowed(scan, tmp_path):
    """A question number on its own must not turn the next image into alt text."""
    pdf = scan(dpi=SCAN_DPI)
    link = reference(pdf, tmp_path)
    markdown = f"12. \n![three chemical structures]({link})\nI\nII\nIII"

    writer = DocxWriter(base_dir=tmp_path)
    render_markdown(writer, markdown)
    writer.save(tmp_path / "numbered-figure.docx")

    document = Document(str(tmp_path / "numbered-figure.docx"))
    assert len(document.inline_shapes) == 1
    assert "![" not in "\n".join(paragraph.text for paragraph in document.paragraphs)


def test_a_reference_out_of_the_working_directory_is_not_embedded(tmp_path):
    """A model that asks for a file outside the job does not get it into the document."""
    secret = tmp_path.parent / "secret.png"
    fitz.Pixmap(fitz.csRGB, fitz.IRect(0, 0, 8, 8)).save(secret)
    (tmp_path / "job").mkdir()

    writer = DocxWriter(base_dir=tmp_path / "job")
    render_markdown(writer, "![a figure](../secret.png)")
    writer.save(tmp_path / "job" / "out.docx")

    assert len(Document(str(tmp_path / "job" / "out.docx")).inline_shapes) == 0


# -- which units the model answered in --------------------------------------- #
#
# The box is asked for on a 0-1000 grid, and models do not reliably give one:
# `anthropic/claude-sonnet-5` — this project's own default — answers in the
# pixels of the image it was shown. Clamping those into the grid used to turn a
# good box into a confidently wrong one, because every coordinate past 1000
# collapsed onto the page's right and bottom edges.

A4_POINTS = (595.0, 842.0)
# The page image a model is shown at the default 180 DPI.
RENDER = (1488, 2105)


def placed(pdf: Path, work_dir: Path, box: str, render=None) -> fitz.Rect | None:
    """The page rectangle a reported box ends up cropping, or None if refused."""
    markdown = f"> _Figure: a diagram_ <!--box: {box}-->"
    result = figures.place_figures(pdf, 0, markdown, work_dir, render=render)
    if not result.startswith("!["):
        return None
    name = result[result.index("(") + 1 : -1]
    png = work_dir / name
    with fitz.open(pdf) as document:
        page = document.load_page(0)
    # Recover the crop's extent on the page from its size and resolution.
    pixmap = fitz.Pixmap(png)
    dpi = pixmap.xres or 96
    return fitz.Rect(0, 0, pixmap.width * 72 / dpi, pixmap.height * 72 / dpi)


def test_a_box_in_pixels_is_rescaled_rather_than_clamped(scan, tmp_path):
    """The same region, named in pixels and on the grid, must crop the same size."""
    pdf = scan(dpi=SCAN_DPI)
    grid = placed(pdf, tmp_path / "grid", "200,600,400,800")
    pixels = placed(
        pdf,
        tmp_path / "px",
        f"{0.2 * RENDER[0]:.0f},{0.6 * RENDER[1]:.0f},{0.4 * RENDER[0]:.0f},{0.8 * RENDER[1]:.0f}",
        render=RENDER,
    )
    assert grid is not None and pixels is not None
    assert pixels.width == pytest.approx(grid.width, rel=0.02)
    assert pixels.height == pytest.approx(grid.height, rel=0.02)


def test_a_pixel_box_inside_the_grid_cannot_be_told_apart(scan, tmp_path):
    """A known and deliberate limit, pinned so it is noticed if it ever changes.

    Pixels only announce themselves by overflowing 1000. A box that sits in the
    top-left corner of the render never does, so it is read on the grid — and no
    amount of arithmetic here can know better from four numbers alone.
    """
    pdf = scan(dpi=SCAN_DPI)
    small = placed(pdf, tmp_path / "a", "200,200,400,400", render=RENDER)
    assert small is not None
    # Read on the grid: a fifth of the page across, not a fifth of the render.
    assert small.width == pytest.approx(200 / 1000 * A4_POINTS[0] + 8.0, abs=2.0)


def test_clamping_a_pixel_box_would_have_cropped_the_wrong_region(scan, tmp_path):
    """Guards the actual defect: the old code kept this box and got it wrong."""
    pdf = scan(dpi=SCAN_DPI)
    # A figure in the lower right of the page, reported in pixels. Clamped to the
    # grid it becomes 1000,1000,1000,1000 — a sliver in the page's bottom corner.
    box = "1015,1580,1165,1745"
    clamped = figures._rect(box, *A4_POINTS)
    assert clamped is None or clamped.width < 5.0, "the old reading was not degenerate"

    rescaled = placed(pdf, tmp_path, box, render=RENDER)
    assert rescaled is not None
    expected_width = (1165 - 1015) / RENDER[0] * A4_POINTS[0]
    assert rescaled.width == pytest.approx(expected_width, abs=12.0)


def test_one_oversized_box_settles_the_units_for_its_whole_page(scan, tmp_path):
    """A page answered in pixels is read in pixels throughout, not box by box."""
    pdf = scan(dpi=SCAN_DPI)
    # Both boxes are pixels; only the second is large enough to prove it.
    markdown = "\n".join(
        [
            "> _Figure: the first_ <!--box: 200,200,400,400-->",
            "> _Figure: the second_ <!--box: 900,1500,1200,1800-->",
        ]
    )
    result = figures.place_figures(pdf, 0, markdown, tmp_path, render=RENDER)
    first = result.splitlines()[0]
    assert first.startswith("![")
    pixmap = fitz.Pixmap(tmp_path / first[first.index("(") + 1 : -1])
    width_pt = pixmap.width * 72 / (pixmap.xres or 96)
    # Read as pixels the box is a fifth of the render, not a fifth of the grid.
    assert width_pt == pytest.approx(200 / RENDER[0] * A4_POINTS[0], abs=12.0)


def test_numbers_that_fit_neither_the_grid_nor_the_render_are_dropped(scan, tmp_path):
    """Better the caption alone than a crop of a region nobody meant."""
    pdf = scan(dpi=SCAN_DPI)
    result = figures.place_figures(
        pdf, 0, "> _Figure: a diagram_ <!--box: 4000,5000,6000,7000-->", tmp_path, render=RENDER
    )
    assert "![" not in result
    assert "a diagram" in result


def test_a_pixel_box_is_dropped_when_the_render_size_is_unknown(scan, tmp_path):
    """Without the render there is nothing to rescale by, and guessing is worse."""
    pdf = scan(dpi=SCAN_DPI)
    result = figures.place_figures(pdf, 0, "> _Figure: a diagram_ <!--box: 1015,1580,1165,1745-->", tmp_path)
    assert "![" not in result
    assert "a diagram" in result


def test_boxes_already_on_the_grid_are_untouched_by_the_render(scan, tmp_path):
    """The common case must not shift because a render size is now passed in."""
    pdf = scan(dpi=SCAN_DPI)
    without = placed(pdf, tmp_path / "a", BOX)
    with_render = placed(pdf, tmp_path / "b", BOX, render=RENDER)
    assert without is not None and with_render is not None
    assert with_render.width == pytest.approx(without.width, rel=0.001)
    assert with_render.height == pytest.approx(without.height, rel=0.001)


def test_a_box_overflowing_the_render_on_one_side_only_is_still_dropped(scan, tmp_path):
    """The case clamping alone does not catch.

    When a single coordinate runs past the render, rescaling leaves the others
    inside the grid and only that one is clamped — which yields a perfectly
    valid rectangle covering most of the page. Nothing later would query it.
    """
    pdf = scan(dpi=SCAN_DPI)
    result = figures.place_figures(
        pdf, 0, "> _Figure: a diagram_ <!--box: 200,100,1600,300-->", tmp_path, render=RENDER
    )
    assert "![" not in result
    assert "a diagram" in result
