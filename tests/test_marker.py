"""The `marker` mode, and the promise it makes.

The promise is that what lands in the document is marker's work. So these tests
are mostly about restraint: that the config reaches marker unaltered including
options this codebase has never heard of, that marker's own output is on disk
untouched, and that the difference between that file and the one the .docx was
built from is only the image prefixes — nothing reordered, nothing repaired.

No sidecar and no network: the client is replaced, so what is under test is this
application's half of the arrangement.
"""

from __future__ import annotations

import base64
import json
from types import SimpleNamespace

import fitz
import pytest
from docx import Document

from app import pipeline
from app.marker_client import MarkerError, parse_document_response

RULE = "-" * 48
PNG_BYTES = (
    b"\x89PNG\r\n\x1a\n\x00\x00\x00\rIHDR\x00\x00\x00\x01\x00\x00\x00\x01\x08\x06\x00\x00\x00"
    b"\x1f\x15\xc4\x89\x00\x00\x00\nIDATx\x9cc\x00\x01\x00\x00\x05\x00\x01\r\n-\xb4\x00\x00"
    b"\x00\x00IEND\xaeB`\x82"
)

MARKDOWN = (
    "# Heat equation\n\n"
    "The diffusion of heat is governed by\n\n"
    "$$\\frac{\\partial u}{\\partial t} = \\alpha \\frac{\\partial^2 u}{\\partial x^2}$$\n\n"
    "![Figure 1](_page_0_Figure_1.png)\n\n"
    "| a | b |\n|---|---|\n| 1 | 2 |\n\n"
    f"{{0}}{RULE}\n\n"
    "## Second page\n\nMore text.\n"
)


class FakeClient:
    """Stands in for the sidecar, and records exactly what it was asked."""

    calls: list[tuple] = []

    def __init__(self):
        FakeClient.calls = []

    def convert_document(self, pdf_path, config=None):
        FakeClient.calls.append((pdf_path, dict(config or {})))
        return parse_document_response({
            "format": "markdown",
            "content": MARKDOWN,
            "images": {"_page_0_Figure_1.png": base64.b64encode(PNG_BYTES).decode("ascii")},
            "metadata": {"pages": 2},
            "config": dict(config or {}),
        })


def settings(**overrides):
    values = {
        "max_pages": 0,
        "marker_options": {},
        "marker_extra_formats": (),
        # Off unless a test is about the page viewer: asking marker for its JSON
        # is a second conversion, and every test here counts the conversions it
        # expects marker to be put through.
        "marker_detection": False,
        "columns": "auto",
    }
    values.update(overrides)
    return SimpleNamespace(**values)


def source_pdf(tmp_path, pages: int = 2):
    path = tmp_path / "source.pdf"
    doc = fitz.open()
    for _ in range(pages):
        doc.new_page(width=595, height=842)
    doc.save(path)
    doc.close()
    return path


def prepare(monkeypatch, **overrides):
    monkeypatch.setattr(pipeline, "settings", settings(**overrides))
    monkeypatch.setattr(pipeline, "MarkerClient", FakeClient)
    # A remote client must never be built in this mode; if anything reaches for
    # one the test should fail loudly rather than quietly try the network.
    monkeypatch.setattr(
        pipeline, "build_client", lambda *a, **k: pytest.fail("marker mode called OpenRouter")
    )


def convert(tmp_path, monkeypatch, columns=None, **overrides):
    prepare(monkeypatch, **overrides)
    pdf = source_pdf(tmp_path)
    work = tmp_path / "job"
    result = pipeline.convert_pdf(
        pdf_path=pdf, work_dir=work, title="Paper", layout="marker", columns=columns
    )
    return result, work


# -- the promise --------------------------------------------------------------- #


def test_markers_own_output_is_kept_verbatim(tmp_path, monkeypatch):
    _, work = convert(tmp_path, monkeypatch)
    assert (work / "marker" / "document.md").read_bytes() == MARKDOWN.encode("utf-8")


def test_the_only_edit_is_the_image_prefix(tmp_path, monkeypatch):
    _, work = convert(tmp_path, monkeypatch)
    built = (work / "document.md").read_text(encoding="utf-8")
    assert built.replace("marker/images/_page_0_Figure_1.png", "_page_0_Figure_1.png") == MARKDOWN


def test_config_reaches_marker_unaltered_including_unknown_options(tmp_path, monkeypatch):
    options = {"use_llm": True, "force_ocr": True, "mode": "fast", "an_option_added_upstream": 7}
    convert(tmp_path, monkeypatch, marker_options=options)
    _, sent = FakeClient.calls[0]
    assert sent == options


def test_max_pages_becomes_a_page_range(tmp_path, monkeypatch):
    convert(tmp_path, monkeypatch, max_pages=3)
    assert FakeClient.calls[0][1]["page_range"] == "0-2"


def test_an_explicit_page_range_wins_over_max_pages(tmp_path, monkeypatch):
    convert(tmp_path, monkeypatch, max_pages=3, marker_options={"page_range": "5-9"})
    assert FakeClient.calls[0][1]["page_range"] == "5-9"


# -- the document that comes out ------------------------------------------------ #


def test_the_mode_costs_nothing_and_makes_no_priced_calls(tmp_path, monkeypatch):
    result, _ = convert(tmp_path, monkeypatch)
    assert result.usage.cost == 0.0
    assert result.usage.calls == 0
    assert [item.extractor for item in result.diagnostics] == ["marker", "marker"]


def test_each_marker_page_becomes_a_word_page(tmp_path, monkeypatch):
    result, work = convert(tmp_path, monkeypatch)
    assert len(result.page_markdown) == 2
    xml = Document(work / "document.docx").element.xml
    assert xml.count('w:type="page"') == 1


def test_equations_arrive_as_native_word_equations(tmp_path, monkeypatch):
    _, work = convert(tmp_path, monkeypatch)
    xml = Document(work / "document.docx").element.xml
    assert "oMathPara" in xml or "<m:oMath" in xml


def test_the_figure_marker_extracted_is_embedded(tmp_path, monkeypatch):
    _, work = convert(tmp_path, monkeypatch)
    assert (work / "marker" / "images" / "_page_0_Figure_1.png").read_bytes() == PNG_BYTES
    document = Document(work / "document.docx")
    assert any(part.content_type.startswith("image/") for part in document.part.package.parts)


def test_metadata_records_what_was_asked_and_what_was_changed(tmp_path, monkeypatch):
    _, work = convert(tmp_path, monkeypatch, marker_options={"mode": "fast"})
    meta = json.loads((work / "marker" / "metadata.json").read_text(encoding="utf-8"))
    assert meta["config"]["mode"] == "fast"
    assert meta["applied"] == {
        "images_prefixed": 1,
        "images_unresolved": 0,
        "pages": 2,
        "paginated": True,
    }


# -- the column choice ---------------------------------------------------------- #
#
# marker returns one linear stream of Markdown whatever the page was set in, so
# the columns are read from the source PDF and the Word section is set to match.
# The choice made in the browser decides whether that happens at all: "natural"
# keeps the one flowing column the transcription already is, "multi" gives the
# source's columns back. Detection has its own tests; what is under test here is
# that the answer reaches the document and is recorded.


def two_column_source(monkeypatch):
    monkeypatch.setattr(pipeline, "detect_columns", lambda *args, **kwargs: [2, 2])


def sections(work):
    return Document(work / "document.docx").element.xml


def test_multi_sets_the_document_in_the_columns_the_source_had(tmp_path, monkeypatch):
    two_column_source(monkeypatch)
    _, work = convert(tmp_path, monkeypatch, columns="multi")
    assert 'w:num="2"' in sections(work)


def test_natural_keeps_one_column_however_the_source_was_set(tmp_path, monkeypatch):
    two_column_source(monkeypatch)
    _, work = convert(tmp_path, monkeypatch, columns="natural")
    assert 'w:num="2"' not in sections(work)
    assert 'w:num="1"' in sections(work)


def test_no_choice_leaves_the_configured_default_in_charge(tmp_path, monkeypatch):
    two_column_source(monkeypatch)
    _, work = convert(tmp_path, monkeypatch, columns=None)  # settings.columns is "auto"
    assert 'w:num="2"' in sections(work)


def test_what_was_asked_for_and_what_it_came_to_are_both_recorded(tmp_path, monkeypatch):
    two_column_source(monkeypatch)
    _, work = convert(tmp_path, monkeypatch, columns="natural")
    meta = json.loads((work / "marker" / "metadata.json").read_text(encoding="utf-8"))
    assert meta["columns_setting"] == "natural"
    assert meta["columns"] == [1, 1]


# -- pages marker read as nothing ----------------------------------------------- #


def blank_page_client(content):
    class Blank(FakeClient):
        def convert_document(self, pdf_path, config=None):
            FakeClient.calls.append((pdf_path, dict(config or {})))
            return parse_document_response({
                "format": "markdown", "content": content, "images": {},
                "metadata": {}, "config": dict(config or {}),
            })

    return Blank


def test_a_page_that_converted_to_nothing_is_reported(tmp_path, monkeypatch):
    """marker calls this a success; the job should not repeat the claim unqualified."""
    content = f"\n\n{{0}}{RULE}\n\nreal text\n\n{{1}}{RULE}\n\n\n\n{{2}}{RULE}\n\nmore text"
    prepare(monkeypatch)
    monkeypatch.setattr(pipeline, "MarkerClient", blank_page_client(content))
    work = tmp_path / "job"
    result = pipeline.convert_pdf(pdf_path=source_pdf(tmp_path), work_dir=work, layout="marker")

    assert [item.fallback_reason for item in result.diagnostics] == [None, "empty_output", None]
    meta = json.loads((work / "marker" / "metadata.json").read_text(encoding="utf-8"))
    assert meta["empty_pages"] == [2]


def test_a_blank_page_still_occupies_its_place_in_the_document(tmp_path, monkeypatch):
    """Dropping it would shift every later page out of step with the PDF."""
    content = f"\n\n{{0}}{RULE}\n\none\n\n{{1}}{RULE}\n\n\n\n{{2}}{RULE}\n\nthree"
    prepare(monkeypatch)
    monkeypatch.setattr(pipeline, "MarkerClient", blank_page_client(content))
    work = tmp_path / "job"
    result = pipeline.convert_pdf(pdf_path=source_pdf(tmp_path), work_dir=work, layout="marker")
    assert len(result.page_markdown) == 3
    assert Document(work / "document.docx").element.xml.count('w:type="page"') == 2


def test_a_wholly_blank_conversion_flags_every_page(tmp_path, monkeypatch):
    """What a degraded inference backend produces: pages, all of them empty."""
    content = f"\n\n{{0}}{RULE}\n\n\n\n{{1}}{RULE}\n\n"
    prepare(monkeypatch)
    monkeypatch.setattr(pipeline, "MarkerClient", blank_page_client(content))
    work = tmp_path / "job"
    result = pipeline.convert_pdf(pdf_path=source_pdf(tmp_path), work_dir=work, layout="marker")
    assert [item.fallback_reason for item in result.diagnostics] == ["empty_output"] * 2


def test_a_page_holding_only_a_figure_is_not_called_empty(tmp_path, monkeypatch):
    content = f"\n\n{{0}}{RULE}\n\ntext\n\n{{1}}{RULE}\n\n![](_page_0_Figure_1.png)"
    prepare(monkeypatch)
    monkeypatch.setattr(pipeline, "MarkerClient", blank_page_client(content))
    work = tmp_path / "job"
    result = pipeline.convert_pdf(pdf_path=source_pdf(tmp_path), work_dir=work, layout="marker")
    assert [item.fallback_reason for item in result.diagnostics] == [None, None]


# -- failure ------------------------------------------------------------------- #


def test_a_sidecar_failure_surfaces_rather_than_producing_a_hollow_document(tmp_path, monkeypatch):
    class Broken(FakeClient):
        def convert_document(self, pdf_path, config=None):
            raise MarkerError("connection refused")

    prepare(monkeypatch)
    monkeypatch.setattr(pipeline, "MarkerClient", Broken)
    work = tmp_path / "job"
    with pytest.raises(MarkerError):
        pipeline.convert_pdf(pdf_path=source_pdf(tmp_path), work_dir=work, layout="marker")
    assert not (work / "document.docx").exists()


def test_an_extra_inspection_format_never_costs_the_job_its_document(tmp_path, monkeypatch):
    class HalfBroken(FakeClient):
        def convert_document(self, pdf_path, config=None):
            if (config or {}).get("output_format") == "json":
                raise MarkerError("that renderer is unavailable")
            return super().convert_document(pdf_path, config)

    prepare(monkeypatch, marker_extra_formats=("json",))
    monkeypatch.setattr(pipeline, "MarkerClient", HalfBroken)
    work = tmp_path / "job"
    pipeline.convert_pdf(pdf_path=source_pdf(tmp_path), work_dir=work, layout="marker")
    assert (work / "document.docx").exists()


# -- the page viewer ----------------------------------------------------------- #

MARKER_JSON = json.dumps({
    "block_type": "Document",
    "metadata": {},
    "children": [{
        "id": "/page/0", "block_type": "Page", "html": "",
        "bbox": [0.0, 0.0, 1000.0, 1414.0], "polygon": [],
        "children": [{
            "id": "/page/0/SectionHeader/1", "block_type": "SectionHeader",
            "html": "<h1>Heat equation</h1>",
            "bbox": [100.0, 120.0, 900.0, 170.0], "polygon": [], "children": None,
        }],
    }],
})


class JsonClient(FakeClient):
    """A sidecar that answers the JSON renderer with real block geometry."""

    def convert_document(self, pdf_path, config=None):
        if (config or {}).get("output_format") == "json":
            FakeClient.calls.append((pdf_path, dict(config or {})))
            return parse_document_response({
                "format": "json", "content": MARKER_JSON, "images": {},
                "metadata": {}, "config": dict(config or {}),
            })
        return super().convert_document(pdf_path, config)


def test_marker_is_asked_for_its_json_so_the_page_viewer_has_boxes(tmp_path, monkeypatch):
    prepare(monkeypatch, marker_detection=True)
    monkeypatch.setattr(pipeline, "MarkerClient", JsonClient)
    work = tmp_path / "job"
    pipeline.convert_pdf(pdf_path=source_pdf(tmp_path), work_dir=work, layout="marker")

    formats = [config.get("output_format") for _, config in FakeClient.calls]
    assert "json" in formats
    detected = json.loads((work / "detection.json").read_text())
    assert detected["mode"] == "marker"
    assert detected["pages"][0]["blocks"][0]["kind"] == "heading"
    # marker's own copy is still the verbatim one it always was.
    assert (work / "marker" / "document.json").read_text() == MARKER_JSON


def test_a_marker_job_still_converts_when_the_boxes_are_turned_off(tmp_path, monkeypatch):
    _, work = convert(tmp_path, monkeypatch, marker_detection=False)

    assert [config.get("output_format") for _, config in FakeClient.calls] == [None]
    assert (work / "document.docx").exists()
    # No geometry, but the pages and their text are still there to read.
    detected = json.loads((work / "detection.json").read_text())
    assert detected["pages"][0]["blocks"] == []
    assert detected["pages"][0]["markdown"]


def test_a_failed_json_pass_never_costs_the_job_its_document(tmp_path, monkeypatch):
    class NoJson(FakeClient):
        def convert_document(self, pdf_path, config=None):
            if (config or {}).get("output_format") == "json":
                raise MarkerError("the JSON renderer fell over")
            return super().convert_document(pdf_path, config)

    prepare(monkeypatch, marker_detection=True)
    monkeypatch.setattr(pipeline, "MarkerClient", NoJson)
    work = tmp_path / "job"
    pipeline.convert_pdf(pdf_path=source_pdf(tmp_path), work_dir=work, layout="marker")

    assert (work / "document.docx").exists()
    assert json.loads((work / "detection.json").read_text())["pages"][0]["blocks"] == []
